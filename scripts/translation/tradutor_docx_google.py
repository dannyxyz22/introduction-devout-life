#!/usr/bin/env python3
"""
Script para gerar arquivo .docx do livro_en.json para tradução no Google Translate.
O Google Translate aceita arquivos .docx de até 10MB para tradução automática.
"""

import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, List, Any

def create_docx_for_translation(input_file: str, output_file: str):
    """
    Cria arquivo .docx estruturado para tradução no Google Translate.
    
    Args:
        input_file (str): Arquivo JSON em inglês
        output_file (str): Arquivo .docx de saída
    """
    print(f"📄 Criando arquivo .docx para tradução...")
    print(f"   Arquivo origem: {input_file}")
    print(f"   Arquivo destino: {output_file}")
    
    # Carrega o arquivo JSON
    with open(input_file, 'r', encoding='utf-8') as f:
        book_data = json.load(f)
    
    # Cria documento Word
    doc = Document()
    
    # Configura título do documento
    title = doc.add_heading('Introduction to the Devout Life - São Francisco de Sales', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph('Documento preparado para tradução no Google Translate').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('Mantenha a estrutura e marcadores ###IDXXXX### para reconstrução automática').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Adiciona quebra de página
    doc.add_page_break()
    
    # Contador para IDs únicos
    id_counter = 1
    
    # Processa cada parte do livro
    for part_idx, part in enumerate(book_data):
        print(f"   📖 Processando Parte {part_idx + 1}...")
        
        # Título da parte
        part_title = part.get('part_title', '')
        if part_title:
            # Adiciona marcador único
            marker = f"###ID{id_counter:04d}###"
            id_counter += 1
            
            # Adiciona título da parte
            part_heading = doc.add_heading(f"PART {part_idx + 1}", level=1)
            part_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Adiciona marcador e texto
            doc.add_paragraph(marker)
            doc.add_paragraph(part_title)
            doc.add_paragraph()  # Linha em branco
        
        # Processa capítulos
        for chapter_idx, chapter in enumerate(part.get('chapters', [])):
            chapter_title = chapter.get('chapter_title', '')
            if chapter_title:
                # Adiciona marcador único
                marker = f"###ID{id_counter:04d}###"
                id_counter += 1
                
                # Adiciona título do capítulo
                chapter_heading = doc.add_heading(f"Chapter {chapter_idx + 1}", level=2)
                
                # Adiciona marcador e texto
                doc.add_paragraph(marker)
                doc.add_paragraph(chapter_title)
                doc.add_paragraph()  # Linha em branco
            
            # Processa conteúdo do capítulo
            for content_idx, content_item in enumerate(chapter.get('content', [])):
                if content_item.get('type') in ['p', 'h1', 'h2', 'h3']:
                    content_text = content_item.get('content', '')
                    if content_text.strip():
                        # Adiciona marcador único
                        marker = f"###ID{id_counter:04d}###"
                        id_counter += 1
                        
                        # Adiciona marcador e conteúdo
                        doc.add_paragraph(marker)
                        
                        # Determina o estilo baseado no tipo
                        if content_item.get('type') == 'h1':
                            doc.add_heading(content_text, level=3)
                        elif content_item.get('type') == 'h2':
                            doc.add_heading(content_text, level=4)
                        elif content_item.get('type') == 'h3':
                            doc.add_heading(content_text, level=5)
                        else:
                            # Parágrafo normal
                            para = doc.add_paragraph(content_text)
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        
                        doc.add_paragraph()  # Linha em branco para separação
    
    # Salva o arquivo
    doc.save(output_file)
    
    # Verifica o tamanho do arquivo
    file_size = os.path.getsize(output_file)
    file_size_mb = file_size / (1024 * 1024)
    
    print(f"\n✅ Arquivo .docx criado com sucesso!")
    print(f"   📂 Arquivo: {output_file}")
    print(f"   📊 Tamanho: {file_size_mb:.2f} MB")
    print(f"   📝 Total de marcadores: {id_counter - 1}")
    
    if file_size_mb > 10:
        print(f"\n⚠️  ATENÇÃO: Arquivo excede 10MB (limite do Google Translate)")
        print(f"   💡 Considere dividir em partes menores")
    else:
        print(f"\n🎉 Arquivo dentro do limite de 10MB do Google Translate!")
    
    return id_counter - 1

def reconstruct_from_docx(docx_file: str, output_json: str, original_json: str):
    """
    Reconstrói o arquivo JSON a partir do .docx traduzido.
    
    Args:
        docx_file (str): Arquivo .docx traduzido
        output_json (str): Arquivo JSON de saída
        original_json (str): Arquivo JSON original para estrutura
    """
    print(f"🔧 Reconstruindo JSON a partir do .docx traduzido...")
    
    try:
        from docx import Document
        doc = Document(docx_file)
    except Exception as e:
        print(f"❌ Erro ao abrir arquivo .docx: {e}")
        return
    
    # Extrai todo o texto do documento
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    # Junta todo o texto
    content = '\n'.join(full_text)
    
    # Extrai traduções usando regex para marcadores
    import re
    translations = {}
    
    # Procura por padrões ###IDXXXX### seguidos de texto
    pattern = r'###ID(\d{4})###\s*\n([^#]*?)(?=###ID\d{4}###|\Z)'
    matches = re.findall(pattern, content, re.DOTALL)
    
    for id_num, text_content in matches:
        marker = f"###ID{id_num}###"
        clean_text = text_content.strip()
        if clean_text:
            translations[marker] = clean_text
    
    print(f"   📊 Traduções extraídas: {len(translations)}")
    
    # Carrega estrutura original
    with open(original_json, 'r', encoding='utf-8') as f:
        book_data = json.load(f)
    
    # Reconstrói com traduções
    translated_book = []
    id_counter = 1
    missing_translations = []
    
    for part_idx, part in enumerate(book_data):
        # Título da parte
        part_title_marker = f"###ID{id_counter:04d}###"
        id_counter += 1
        
        translated_part_title = translations.get(part_title_marker, part.get('part_title', ''))
        if part_title_marker not in translations and part.get('part_title'):
            missing_translations.append(part_title_marker)
        
        translated_part = {
            "part_title": translated_part_title,
            "chapters": []
        }
        
        # Processa capítulos
        for chapter_idx, chapter in enumerate(part.get('chapters', [])):
            chapter_title_marker = f"###ID{id_counter:04d}###"
            id_counter += 1
            
            translated_chapter_title = translations.get(chapter_title_marker, chapter.get('chapter_title', ''))
            if chapter_title_marker not in translations and chapter.get('chapter_title'):
                missing_translations.append(chapter_title_marker)
            
            translated_chapter = {
                "chapter_title": translated_chapter_title,
                "content": []
            }
            
            # Processa conteúdo
            for content_idx, content_item in enumerate(chapter.get('content', [])):
                if content_item.get('type') in ['p', 'h1', 'h2', 'h3']:
                    content_marker = f"###ID{id_counter:04d}###"
                    id_counter += 1
                    
                    translated_content = translations.get(content_marker, content_item.get('content', ''))
                    if content_marker not in translations and content_item.get('content', '').strip():
                        missing_translations.append(content_marker)
                    
                    translated_item = {
                        'type': content_item.get('type'),
                        'content': translated_content,
                        'word_count': len(translated_content.split()) if translated_content else 0
                    }
                else:
                    translated_item = content_item.copy()
                
                translated_chapter['content'].append(translated_item)
            
            translated_part['chapters'].append(translated_chapter)
        
        translated_book.append(translated_part)
    
    # Salva arquivo JSON traduzido
    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(translated_book, f, ensure_ascii=False, indent=2)
    
    print(f"✅ Reconstrução concluída!")
    print(f"   📂 Arquivo salvo: {output_json}")
    
    if missing_translations:
        print(f"\n⚠️  {len(missing_translations)} traduções não encontradas (mantido texto original)")
    else:
        print(f"\n🎉 Todas as traduções foram aplicadas com sucesso!")

def main():
    """Função principal"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Arquivos
    input_json = os.path.join(script_dir, 'leitura-devota-app', 'public', 'data', 'livro_en.json')
    output_docx = os.path.join(script_dir, 'livro_para_traducao.docx')
    output_json = os.path.join(script_dir, 'leitura-devota-app', 'public', 'data', 'livro_pt-BR.json')
    
    print("📄 TRADUTOR DOCX - GOOGLE TRANSLATE")
    print("=" * 50)
    print("1. Gerar arquivo .docx para tradução")
    print("2. Reconstruir JSON a partir do .docx traduzido")
    print("3. Sair")
    
    choice = input("\nEscolha uma opção (1-3): ").strip()
    
    if choice == '1':
        if not os.path.exists(input_json):
            print(f"❌ Erro: Arquivo não encontrado: {input_json}")
            return
        
        total_markers = create_docx_for_translation(input_json, output_docx)
        
        print(f"\n📝 INSTRUÇÕES PARA TRADUÇÃO:")
        print(f"   1. Abra o Google Translate: https://translate.google.com")
        print(f"   2. Configure: Inglês → Português")
        print(f"   3. Clique em 'Documentos'")
        print(f"   4. Faça upload do arquivo: {output_docx}")
        print(f"   5. Aguarde a tradução automática")
        print(f"   6. Baixe o arquivo traduzido")
        print(f"   7. Execute a opção 2 deste script com o arquivo baixado")
        
    elif choice == '2':
        translated_docx = input("Digite o caminho do arquivo .docx traduzido: ").strip()
        
        if not os.path.exists(translated_docx):
            print(f"❌ Erro: Arquivo não encontrado: {translated_docx}")
            return
        
        if not os.path.exists(input_json):
            print(f"❌ Erro: Arquivo original não encontrado: {input_json}")
            return
        
        # Cria backup se necessário
        if os.path.exists(output_json):
            backup_file = output_json.replace('.json', '_backup.json')
            print(f"📋 Criando backup: {backup_file}")
            
            with open(output_json, 'r', encoding='utf-8') as f:
                data = f.read()
            
            with open(backup_file, 'w', encoding='utf-8') as f:
                f.write(data)
        
        reconstruct_from_docx(translated_docx, output_json, input_json)
        
    elif choice == '3':
        print("👋 Até logo!")
        
    else:
        print("❌ Opção inválida!")

if __name__ == '__main__':
    main()
