#!/usr/bin/env python3
"""
Script CORRIGIDO para gerar arquivo .docx APENAS com conteúdo textual puro.
Remove metadados como "Chapter 1", "Part 1" etc. que contaminam a tradução.
Inclui também a Oração Dedicatória e o Prefácio.
"""

import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, List, Any
import xml.etree.ElementTree as ET

def extract_text_from_xhtml(xhtml_file: str) -> List[str]:
    """
    Extrai texto de um arquivo XHTML, preservando parágrafos.
    Função reutilizável seguindo o princípio DRY.
    
    Args:
        xhtml_file (str): Caminho para o arquivo XHTML
        
    Returns:
        List[str]: Lista de parágrafos extraídos
    """
    if not os.path.exists(xhtml_file):
        return []
    
    try:
        tree = ET.parse(xhtml_file)
        root = tree.getroot()
        
        # Namespace XHTML
        ns = {'xhtml': 'http://www.w3.org/1999/xhtml'}
        
        paragraphs = []
        
        # Extrair título h1
        h1_elements = root.findall('.//xhtml:h1', ns)
        for h1 in h1_elements:
            if h1.text and h1.text.strip():
                paragraphs.append(h1.text.strip())
        
        # Extrair parágrafos
        p_elements = root.findall('.//xhtml:p', ns)
        for p in p_elements:
            # Concatenar todo o texto do parágrafo (incluindo de elementos filhos)
            text_parts = []
            if p.text:
                text_parts.append(p.text)
            
            for child in p:
                if child.text:
                    text_parts.append(child.text)
                if child.tail:
                    text_parts.append(child.tail)
            
            full_text = ''.join(text_parts).strip()
            if full_text and not full_text.startswith('<!--'):  # Ignora comentários
                paragraphs.append(full_text)
        
        return paragraphs
    except Exception as e:
        print(f"   ⚠️ Erro ao processar {xhtml_file}: {e}")
        return []

def add_xhtml_content_to_docx(doc: Document, xhtml_file: str, content_type: str, id_counter: int) -> int:
    """
    Adiciona conteúdo de um arquivo XHTML ao documento DOCX.
    Função reutilizável seguindo o princípio DRY.
    
    Args:
        doc (Document): Documento DOCX
        xhtml_file (str): Caminho para o arquivo XHTML
        content_type (str): Tipo de conteúdo para logging
        id_counter (int): Contador de IDs atual
        
    Returns:
        int: Novo valor do contador de IDs
    """
    paragraphs = extract_text_from_xhtml(xhtml_file)
    
    if paragraphs:
        print(f"   📄 Incluindo {content_type}...")
        for paragraph in paragraphs:
            marker = f"###ID{id_counter:04d}###"
            id_counter += 1
            
            doc.add_paragraph(marker)
            para = doc.add_paragraph(paragraph)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            doc.add_paragraph()  # Linha em branco
    
    return id_counter

def create_clean_docx_for_translation(input_file: str, output_file: str):
    """
    Cria arquivo .docx LIMPO com APENAS conteúdo textual para tradução.
    Remove todos os metadados que podem contaminar a tradução automática.
    Inclui Oração Dedicatória e Prefácio seguindo o princípio DRY.
    
    Args:
        input_file (str): Arquivo JSON em inglês
        output_file (str): Arquivo .docx de saída
    """
    print(f"🧹 Criando arquivo .docx LIMPO para tradução...")
    print(f"   ℹ️  Incluindo Oração Dedicatória e Prefácio")
    print(f"   ℹ️  Removendo metadados como 'Chapter 1', 'Part 1' etc.")
    print(f"   📂 Origem: {input_file}")
    print(f"   📂 Destino: {output_file}")
    
    # Carrega o arquivo JSON
    with open(input_file, 'r', encoding='utf-8') as f:
        book_data = json.load(f)
    
    # Cria documento Word minimalista
    doc = Document()
    
    # APENAS um título simples (que pode ser traduzido)
    title = doc.add_heading('Introduction to the Devout Life', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Informações técnicas em comentário (não serão traduzidas)
    doc.add_paragraph('<!-- TECHNICAL INFO: Keep ###IDXXXX### markers for reconstruction -->').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()
    
    # Contador para IDs únicos
    id_counter = 1
    total_texts = 0
    
    # Determinar caminhos dos arquivos XHTML (princípio DRY)
    script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # Volta 2 níveis
    epub_processing_dir = os.path.join(script_dir, 'epub_processing')
    
    # Incluir Oração Dedicatória (reutilizando função)
    dedicatory_prayer_file = os.path.join(epub_processing_dir, 'dedicatory_prayer_en.xhtml')
    old_counter = id_counter
    id_counter = add_xhtml_content_to_docx(doc, dedicatory_prayer_file, "Oração Dedicatória", id_counter)
    total_texts += (id_counter - old_counter)
    
    # Incluir Prefácio (reutilizando função)
    preface_file = os.path.join(epub_processing_dir, 'preface_en.xhtml')
    old_counter = id_counter
    id_counter = add_xhtml_content_to_docx(doc, preface_file, "Prefácio", id_counter)
    total_texts += (id_counter - old_counter)
    
    # Processa cada parte do livro
    for part_idx, part in enumerate(book_data):
        print(f"   📖 Processando Parte {part_idx + 1}...")
        
        # APENAS o conteúdo do título da parte (SEM "PART 1")
        part_title = part.get('part_title', '')
        if part_title:
            marker = f"###ID{id_counter:04d}###"
            id_counter += 1
            total_texts += 1
            
            # APENAS o texto, sem marcadores de "PART X"
            doc.add_paragraph(marker)
            para = doc.add_paragraph(part_title)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            doc.add_paragraph()  # Linha em branco
        
        # Processa capítulos
        for chapter_idx, chapter in enumerate(part.get('chapters', [])):
            chapter_title = chapter.get('chapter_title', '')
            if chapter_title:
                marker = f"###ID{id_counter:04d}###"
                id_counter += 1
                total_texts += 1
                
                # APENAS o texto do título, sem "Chapter X"
                doc.add_paragraph(marker)
                para = doc.add_paragraph(chapter_title)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                doc.add_paragraph()  # Linha em branco
            
            # Processa conteúdo do capítulo
            for content_item in chapter.get('content', []):
                if content_item.get('type') in ['p', 'h1', 'h2', 'h3']:
                    content_text = content_item.get('content', '')
                    if content_text.strip():
                        marker = f"###ID{id_counter:04d}###"
                        id_counter += 1
                        total_texts += 1
                        
                        # APENAS o conteúdo textual puro
                        doc.add_paragraph(marker)
                        para = doc.add_paragraph(content_text)
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        doc.add_paragraph()  # Linha em branco
    
    # Salva o arquivo
    doc.save(output_file)
    
    # Verifica o tamanho do arquivo
    file_size = os.path.getsize(output_file)
    file_size_mb = file_size / (1024 * 1024)
    
    print(f"\n✅ Arquivo .docx LIMPO criado com sucesso!")
    print(f"   📂 Arquivo: {output_file}")
    print(f"   📊 Tamanho: {file_size_mb:.2f} MB")
    print(f"   📝 Total de textos: {total_texts}")
    print(f"   🎯 SEM metadados contaminantes!")
    
    if file_size_mb > 10:
        print(f"   ⚠️  ATENÇÃO: Arquivo maior que 10MB. Google Translate tem limite de 10MB.")
        return False
    
    return True

def create_translated_xhtml_files(translated_texts: dict):
    """
    Cria arquivos XHTML traduzidos a partir dos textos do DOCX.
    Segue o princípio DRY reutilizando a estrutura do gerador EPUB.
    
    Args:
        translated_texts (dict): Dicionário com textos traduzidos {id: texto}
    """
    print("   📝 Criando arquivos XHTML traduzidos...")
    
    script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    epub_processing_dir = os.path.join(script_dir, 'epub_processing')
    
    # Cria oração dedicatória traduzida
    dedicatory_id = "###ID1###"
    if dedicatory_id in translated_texts:
        dedicatory_content = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>Oração Dedicatória</title>
    <meta http-equiv="Content-Type" content="text/html; charset=utf-8"/>
    <style type="text/css">
        body {{
            font-family: serif;
            margin: 2em;
            line-height: 1.6;
            background: transparent;
            color: currentColor;
        }}
        .dedicatory {{
            text-align: center;
            padding: 2em 1em;
            margin: 2em auto;
            max-width: 600px;
            border: 2px solid currentColor;
            background: transparent;
        }}
        .prayer-text {{
            font-style: italic;
            margin: 1em 0;
            color: currentColor;
        }}
    </style>
</head>
<body>
    <div class="dedicatory">
        <h2>Oração Dedicatória</h2>
        <div class="prayer-text">
            {translated_texts[dedicatory_id]}
        </div>
    </div>
</body>
</html>"""
        
        dedicatory_file = os.path.join(epub_processing_dir, 'dedicatory_prayer_pt-BR.xhtml')
        with open(dedicatory_file, 'w', encoding='utf-8') as f:
            f.write(dedicatory_content)
        print(f"   ✅ Oração dedicatória criada: {dedicatory_file}")
    
    # Cria prefácio traduzido
    preface_id = "###ID2###"  # Assumindo que o prefácio é o segundo texto
    if preface_id in translated_texts:
        preface_content = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>Prefácio</title>
    <meta http-equiv="Content-Type" content="text/html; charset=utf-8"/>
    <style type="text/css">
        body {{
            font-family: serif;
            margin: 2em;
            line-height: 1.6;
            background: transparent;
            color: currentColor;
        }}
        .preface {{
            text-align: justify;
            max-width: 700px;
            margin: 0 auto;
        }}
        .preface p {{
            margin: 1em 0;
            text-indent: 2em;
        }}
    </style>
</head>
<body>
    <div class="preface">
        <h2>Prefácio</h2>
        <p>{translated_texts[preface_id]}</p>
    </div>
</body>
</html>"""
        
        preface_file = os.path.join(epub_processing_dir, 'preface_pt-BR.xhtml')
        with open(preface_file, 'w', encoding='utf-8') as f:
            f.write(preface_content)
        print(f"   ✅ Prefácio criado: {preface_file}")


def reconstruct_from_clean_docx(docx_file: str, output_json: str, original_json: str):
    """
    Reconstrói o arquivo JSON a partir do .docx traduzido LIMPO.
    Inclui processamento da Oração Dedicatória e Prefácio traduzidos.
    
    Args:
        docx_file (str): Arquivo .docx traduzido pelo Google Translate
        output_json (str): Arquivo JSON de saída em português
        original_json (str): Arquivo JSON original em inglês (para estrutura)
    """
    print(f"🔄 Reconstruindo JSON a partir do .docx traduzido...")
    print(f"   📂 Arquivo traduzido: {docx_file}")
    print(f"   📂 JSON original: {original_json}")
    print(f"   📂 JSON de saída: {output_json}")
    
    # Carrega arquivo original para manter estrutura
    with open(original_json, 'r', encoding='utf-8') as f:
        original_data = json.load(f)
    
    # Extrai textos traduzidos do .docx
    doc = Document(docx_file)
    translated_texts = {}
    
    # Processa parágrafos do documento
    current_id = None
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Verifica se é um marcador ID
        if text.startswith('###ID') and text.endswith('###'):
            current_id = text
        elif current_id and text and not text.startswith('<!--'):
            # É um texto traduzido
            translated_texts[current_id] = text
            current_id = None
    
    print(f"   📝 Textos traduzidos extraídos: {len(translated_texts)}")
    
    # Cria arquivos XHTML traduzidos (princípio DRY - reutiliza estrutura do gerador EPUB)
    create_translated_xhtml_files(translated_texts)
    
    # Reconstrói estrutura com textos traduzidos
    id_counter = 1
    
    # Pula IDs usados pela oração dedicatória e prefácio
    script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    epub_processing_dir = os.path.join(script_dir, 'epub_processing')
    
    # Conta textos da oração dedicatória
    dedicatory_prayer_file = os.path.join(epub_processing_dir, 'dedicatory_prayer_en.xhtml')
    dedicatory_count = len(extract_text_from_xhtml(dedicatory_prayer_file))
    id_counter += dedicatory_count
    
    # Conta textos do prefácio
    preface_file = os.path.join(epub_processing_dir, 'preface_en.xhtml')
    preface_count = len(extract_text_from_xhtml(preface_file))
    id_counter += preface_count
    
    for part in original_data:
        # Traduz título da parte
        if 'part_title' in part:
            marker = f"###ID{id_counter:04d}###"
            if marker in translated_texts:
                part['part_title'] = translated_texts[marker]
            id_counter += 1
        
        # Traduz capítulos
        for chapter in part.get('chapters', []):
            # Traduz título do capítulo
            if 'chapter_title' in chapter:
                marker = f"###ID{id_counter:04d}###"
                if marker in translated_texts:
                    chapter['chapter_title'] = translated_texts[marker]
                id_counter += 1
            
            # Traduz conteúdo
            for content_item in chapter.get('content', []):
                if content_item.get('type') in ['p', 'h1', 'h2', 'h3'] and content_item.get('content', '').strip():
                    marker = f"###ID{id_counter:04d}###"
                    if marker in translated_texts:
                        content_item['content'] = translated_texts[marker]
                        # Recalcula word count
                        content_item['word_count'] = len(translated_texts[marker].split())
                    id_counter += 1
    
    # Salva arquivo JSON traduzido
    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(original_data, f, indent=2, ensure_ascii=False)
    
    print(f"✅ Arquivo JSON em português criado: {output_json}")
    
    # Estatísticas
    total_parts = len(original_data)
    total_chapters = sum(len(part.get('chapters', [])) for part in original_data)
    total_content = sum(
        len(chapter.get('content', [])) 
        for part in original_data 
        for chapter in part.get('chapters', [])
    )
    
    print(f"\n📊 ESTATÍSTICAS DA TRADUÇÃO:")
    print(f"   📚 Partes: {total_parts}")
    print(f"   📖 Capítulos: {total_chapters}")
    print(f"   📝 Itens de conteúdo: {total_content}")
    print(f"   🔄 Textos traduzidos aplicados: {len(translated_texts)}")

def main():
    """Função principal"""
    print("🧹 GERADOR DE DOCX LIMPO PARA TRADUÇÃO")
    print("Remove metadados contaminantes como 'Chapter 1', 'Part 1' etc.")
    print("=" * 60)
    
    # Arquivos
    # Detectar diretório base do projeto
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(os.path.dirname(script_dir))
    input_json = os.path.join(project_root, 'webapp', 'public', 'data', 'livro_en.json')
    output_docx = 'livro_en_CLEAN_for_translation.docx'
    
    if not os.path.exists(input_json):
        print(f"❌ Arquivo não encontrado: {input_json}")
        return
    
    # Gera .docx limpo
    success = create_clean_docx_for_translation(input_json, output_docx)
    
    if success:
        print(f"\n🎯 PRÓXIMOS PASSOS:")
        print(f"1. 📤 Faça upload do arquivo '{output_docx}' no Google Translate")
        print(f"2. 🇧🇷 Traduza de Inglês para Português")
        print(f"3. 📥 Baixe o arquivo traduzido")
        print(f"4. 🔄 Execute a função de reconstrução para gerar o JSON em português")
        print(f"\n✨ O arquivo está LIMPO, sem metadados contaminantes!")
    else:
        print(f"\n❌ Erro na criação do arquivo .docx")

if __name__ == "__main__":
    main()
